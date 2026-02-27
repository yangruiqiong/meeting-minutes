import streamlit as st
import datetime
import os
import io
import numpy as np
import tempfile

st.set_page_config(page_title="会议纪要助手", page_icon="🎙️", layout="centered")

# ── Supabase 初始化 ──────────────────────────────────────────────
def get_supabase():
    try:
        from supabase import create_client
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except Exception:
        return None

def save_to_db(supabase, topic, attendees, content, source):
    if not supabase:
        return False
    try:
        supabase.table("meeting_minutes").insert({
            "topic": topic or "（未填写）",
            "attendees": attendees or "（未填写）",
            "content": content,
            "source": source,
            "created_at": datetime.datetime.now().isoformat()
        }).execute()
        return True
    except Exception as e:
        st.warning(f"保存历史记录失败：{e}")
        return False

def load_history(supabase):
    if not supabase:
        return []
    try:
        res = supabase.table("meeting_minutes").select("*").order("created_at", desc=True).execute()
        return res.data
    except Exception:
        return []

def delete_record(supabase, record_id):
    if not supabase:
        return
    try:
        supabase.table("meeting_minutes").delete().eq("id", record_id).execute()
    except Exception as e:
        st.warning(f"删除失败：{e}")

# ── 语音转文字 ───────────────────────────────────────────────────
@st.cache_resource(show_spinner="正在加载语音模型...")
def load_whisper():
    import whisper
    return whisper.load_model("base")

def transcribe_audio_bytes(audio_bytes, suffix=".wav"):
    import soundfile as sf
    model = load_whisper()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(audio_bytes)
        tmp_path = tmp.name
    try:
        data, sr = sf.read(tmp_path, dtype='float32')
        if data.ndim > 1:
            data = data.mean(axis=1)
        if sr != 16000:
            new_len = int(len(data) / sr * 16000)
            data = np.interp(np.linspace(0, len(data), new_len),
                             np.arange(len(data)), data).astype('float32')
        result = model.transcribe(data, language='zh')
        return result['text'].strip()
    finally:
        os.unlink(tmp_path)

# ── 生成纪要 ────────────────────────────────────────────────────
def generate_minutes(content, attendees, topic, source):
    now = datetime.datetime.now()
    date_str = now.strftime("%Y年%m月%d日 %H:%M")
    sentences = [s.strip() for s in
                 content.replace('。', '。\n').replace('！', '！\n').replace('？', '？\n').split('\n')
                 if s.strip()]
    minutes = f"""会议纪要
{'='*40}
会议时间：{date_str}
参会人员：{attendees or '（未填写）'}
会议主题：{topic or '（未填写）'}
来　　源：{source}

【原始内容】
{content}

【要点整理】
"""
    for i, s in enumerate(sentences[:15], 1):
        minutes += f"{i}. {s}\n"
    minutes += f"\n【待办事项】\n（请手动补充）\n\n{'='*40}\n生成时间：{date_str}\n"
    return minutes

def minutes_to_docx(text):
    from docx import Document as DocxDoc
    doc = DocxDoc()
    doc.add_heading('会议纪要', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════
# 主界面
# ═══════════════════════════════════════════════════════════════
supabase = get_supabase()

page = st.sidebar.radio("📌 导航", ["✍️ 新建会议纪要", "📚 历史记录"])

# ──────────────────────────────────────────────────────────────
# 页面一：新建会议纪要
# ──────────────────────────────────────────────────────────────
if page == "✍️ 新建会议纪要":
    st.title("🎙️ 会议纪要助手")
    st.caption("录音 / 上传音频 / 上传手写记录 → 自动生成会议纪要")

    col1, col2 = st.columns(2)
    with col1:
        attendees = st.text_input("👥 参会人员", placeholder="张三、李四、王五")
    with col2:
        topic = st.text_input("📌 会议主题", placeholder="产品上线讨论")

    st.divider()

    tab1, tab2, tab3 = st.tabs(["🎙️ 录音", "📂 上传音频", "📄 上传手写记录"])
    transcript = None
    source = ""

    with tab1:
        st.info("点击麦克风按钮开始录音，录完后点停止")
        audio_value = st.audio_input("录音")
        if audio_value and st.button("🔄 转文字并生成会议纪要", key="btn_record"):
            with st.spinner("正在识别语音..."):
                transcript = transcribe_audio_bytes(audio_value.getvalue())
                source = "语音录音转写"

    with tab2:
        audio_file = st.file_uploader("上传音频", type=["wav", "mp3", "m4a", "ogg", "flac"])
        if audio_file:
            st.audio(audio_file)
            if st.button("🔄 转文字并生成会议纪要", key="btn_audio"):
                with st.spinner("正在识别语音..."):
                    ext = "." + audio_file.name.split(".")[-1]
                    transcript = transcribe_audio_bytes(audio_file.getvalue(), suffix=ext)
                    source = "音频文件转写"

    with tab3:
        st.info("上传 Word (.docx) 或文本 (.txt) 文件，自动整理为标准会议纪要")
        doc_file = st.file_uploader("上传文件", type=["docx", "txt"])
        if doc_file and st.button("📋 整理为会议纪要", key="btn_doc"):
            if doc_file.name.endswith(".docx"):
                from docx import Document
                doc = Document(io.BytesIO(doc_file.getvalue()))
                transcript = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                transcript = doc_file.getvalue().decode('utf-8')
            source = "手写记录整理"

    if transcript:
        st.divider()
        st.subheader("📝 会议纪要")
        minutes = generate_minutes(transcript, attendees, topic, source)
        now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M")

        edited = st.text_area("会议纪要（可直接编辑）", value=minutes, height=400)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("💾 下载 TXT", data=edited.encode('utf-8'),
                               file_name=f"会议纪要_{now_str}.txt", mime="text/plain")
        with col2:
            st.download_button("📝 下载 Word", data=minutes_to_docx(edited),
                               file_name=f"会议纪要_{now_str}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col3:
            if st.button("💿 保存到历史记录"):
                if supabase:
                    ok = save_to_db(supabase, topic, attendees, edited, source)
                    if ok:
                        st.success("✅ 已保存到历史记录！")
                else:
                    st.error("未配置数据库，无法保存历史记录")

# ──────────────────────────────────────────────────────────────
# 页面二：历史记录
# ──────────────────────────────────────────────────────────────
elif page == "📚 历史记录":
    st.title("📚 历史会议记录")

    if not supabase:
        st.error("未配置数据库连接，请先在 Streamlit Secrets 中配置 SUPABASE_URL 和 SUPABASE_KEY")
        st.stop()

    # 搜索
    keyword = st.text_input("🔍 搜索（主题 / 参会人 / 内容）", placeholder="输入关键词")

    records = load_history(supabase)

    if keyword:
        records = [r for r in records if
                   keyword.lower() in (r.get('topic') or '').lower() or
                   keyword.lower() in (r.get('attendees') or '').lower() or
                   keyword.lower() in (r.get('content') or '').lower()]

    if not records:
        st.info("暂无历史记录" if not keyword else "没有找到相关记录")
    else:
        st.caption(f"共 {len(records)} 条记录")
        for r in records:
            created = r.get('created_at', '')[:16].replace('T', ' ')
            with st.expander(f"📄 {r.get('topic', '无主题')}  |  {r.get('attendees', '')}  |  {created}"):
                st.text_area("内容", value=r.get('content', ''), height=300,
                             key=f"content_{r['id']}", disabled=False)
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("💾 下载 TXT",
                                       data=(r.get('content') or '').encode('utf-8'),
                                       file_name=f"会议纪要_{created.replace(' ','_')}.txt",
                                       key=f"dl_txt_{r['id']}")
                with col2:
                    st.download_button("📝 下载 Word",
                                       data=minutes_to_docx(r.get('content') or ''),
                                       file_name=f"会议纪要_{created.replace(' ','_')}.docx",
                                       key=f"dl_docx_{r['id']}")
                with col3:
                    if st.button("🗑️ 删除", key=f"del_{r['id']}"):
                        delete_record(supabase, r['id'])
                        st.rerun()
