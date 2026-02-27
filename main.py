import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import os
import datetime

# 检查并安装依赖
def check_dependencies():
    import subprocess, sys
    packages = {
        'sounddevice': 'sounddevice',
        'soundfile': 'soundfile',
        'numpy': 'numpy',
        'whisper': 'openai-whisper',
        'docx': 'python-docx',
    }
    mirror = '-i https://pypi.tuna.tsinghua.edu.cn/simple --trusted-host pypi.tuna.tsinghua.edu.cn'
    for module, package in packages.items():
        try:
            __import__(module)
        except ImportError:
            print(f"正在安装 {package}...")
            subprocess.check_call(f'{sys.executable} -m pip install {package} {mirror}', shell=True)

check_dependencies()

import sounddevice as sd
import soundfile as sf
import numpy as np
import whisper
from docx import Document

class MeetingMinutesApp:
    def __init__(self, root):
        self.root = root
        self.root.title("会议纪要助手")
        self.root.geometry("700x680")
        self.root.configure(bg='#f0f0f0')

        self.recording = False
        self.audio_data = []
        self.audio_array = None
        self.sample_rate = 16000
        self.model = None
        self.audio_file = None

        self.build_ui()

    def build_ui(self):
        title = tk.Label(self.root, text="🎙️ 会议纪要助手", font=('微软雅黑', 18, 'bold'),
                         bg='#f0f0f0', fg='#333')
        title.pack(pady=15)

        # 会议信息
        frame_info = tk.LabelFrame(self.root, text="会议信息", font=('微软雅黑', 10),
                                    bg='#f0f0f0', padx=10, pady=8)
        frame_info.pack(fill='x', padx=20, pady=5)

        tk.Label(frame_info, text="参会人员：", bg='#f0f0f0', font=('微软雅黑', 10)).grid(row=0, column=0, sticky='w')
        self.attendees_var = tk.StringVar()
        tk.Entry(frame_info, textvariable=self.attendees_var, width=50, font=('微软雅黑', 10)).grid(row=0, column=1, padx=5)

        tk.Label(frame_info, text="会议主题：", bg='#f0f0f0', font=('微软雅黑', 10)).grid(row=1, column=0, sticky='w', pady=5)
        self.topic_var = tk.StringVar()
        tk.Entry(frame_info, textvariable=self.topic_var, width=50, font=('微软雅黑', 10)).grid(row=1, column=1, padx=5)

        # 录音控制
        frame_record = tk.LabelFrame(self.root, text="方式一：录音", font=('微软雅黑', 10),
                                      bg='#f0f0f0', padx=10, pady=8)
        frame_record.pack(fill='x', padx=20, pady=5)

        btn_frame = tk.Frame(frame_record, bg='#f0f0f0')
        btn_frame.pack()

        self.btn_start = tk.Button(btn_frame, text="▶ 开始录音", command=self.start_recording,
                                    bg='#4CAF50', fg='white', font=('微软雅黑', 11, 'bold'),
                                    padx=20, pady=8, relief='flat', cursor='hand2')
        self.btn_start.grid(row=0, column=0, padx=8)

        self.btn_stop = tk.Button(btn_frame, text="⏹ 停止录音", command=self.stop_recording,
                                   bg='#f44336', fg='white', font=('微软雅黑', 11, 'bold'),
                                   padx=20, pady=8, relief='flat', cursor='hand2', state='disabled')
        self.btn_stop.grid(row=0, column=1, padx=8)

        self.btn_import_audio = tk.Button(btn_frame, text="📂 导入音频", command=self.import_audio,
                                     bg='#2196F3', fg='white', font=('微软雅黑', 11, 'bold'),
                                     padx=20, pady=8, relief='flat', cursor='hand2')
        self.btn_import_audio.grid(row=0, column=2, padx=8)

        self.status_label = tk.Label(frame_record, text="状态：待机", bg='#f0f0f0',
                                      font=('微软雅黑', 10), fg='#666')
        self.status_label.pack(pady=5)

        self.btn_transcribe = tk.Button(self.root, text="🔄 转文字并生成会议纪要",
                                         command=self.transcribe_and_generate,
                                         bg='#FF9800', fg='white', font=('微软雅黑', 12, 'bold'),
                                         padx=30, pady=10, relief='flat', cursor='hand2', state='disabled')
        self.btn_transcribe.pack(pady=5)

        # Word 导入
        frame_word = tk.LabelFrame(self.root, text="方式二：导入手写记录（Word/TXT）", font=('微软雅黑', 10),
                                    bg='#f0f0f0', padx=10, pady=8)
        frame_word.pack(fill='x', padx=20, pady=5)

        self.btn_import_word = tk.Button(frame_word, text="📄 导入 Word / TXT 文件，整理为会议纪要",
                                          command=self.import_word,
                                          bg='#607D8B', fg='white', font=('微软雅黑', 11, 'bold'),
                                          padx=20, pady=8, relief='flat', cursor='hand2')
        self.btn_import_word.pack()

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='indeterminate', length=400)
        self.progress.pack(pady=5)

        # 结果显示
        frame_result = tk.LabelFrame(self.root, text="会议纪要", font=('微软雅黑', 10),
                                      bg='#f0f0f0', padx=10, pady=8)
        frame_result.pack(fill='both', expand=True, padx=20, pady=5)

        self.result_text = tk.Text(frame_result, font=('微软雅黑', 10), wrap='word',
                                    bg='white', relief='flat', padx=8, pady=8)
        scrollbar = ttk.Scrollbar(frame_result, command=self.result_text.yview)
        self.result_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        self.result_text.pack(fill='both', expand=True)

        # 保存按钮
        save_frame = tk.Frame(self.root, bg='#f0f0f0')
        save_frame.pack(pady=10)

        self.btn_save_txt = tk.Button(save_frame, text="💾 保存为 TXT", command=lambda: self.save_minutes('txt'),
                                       bg='#9C27B0', fg='white', font=('微软雅黑', 11, 'bold'),
                                       padx=20, pady=8, relief='flat', cursor='hand2', state='disabled')
        self.btn_save_txt.grid(row=0, column=0, padx=10)

        self.btn_save_word = tk.Button(save_frame, text="📝 保存为 Word", command=lambda: self.save_minutes('docx'),
                                        bg='#1565C0', fg='white', font=('微软雅黑', 11, 'bold'),
                                        padx=20, pady=8, relief='flat', cursor='hand2', state='disabled')
        self.btn_save_word.grid(row=0, column=1, padx=10)

    def start_recording(self):
        self.recording = True
        self.audio_data = []
        self.audio_array = None
        self.btn_start.config(state='disabled')
        self.btn_stop.config(state='normal')
        self.btn_import_audio.config(state='disabled')
        self.status_label.config(text="状态：🔴 录音中...", fg='red')

        def record():
            with sd.InputStream(samplerate=self.sample_rate, channels=1, dtype='float32') as stream:
                while self.recording:
                    data, _ = stream.read(1024)
                    self.audio_data.append(data.copy())

        self.record_thread = threading.Thread(target=record, daemon=True)
        self.record_thread.start()

    def stop_recording(self):
        self.recording = False
        self.btn_start.config(state='normal')
        self.btn_stop.config(state='disabled')
        self.btn_import_audio.config(state='normal')

        if self.audio_data:
            audio_array = np.concatenate(self.audio_data, axis=0)
            self.audio_array = audio_array
            self.audio_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'temp_recording.wav')
            sf.write(self.audio_file, audio_array, self.sample_rate)
            self.btn_transcribe.config(state='normal')
            self.status_label.config(text=f"状态：✅ 录音已保存，时长约 {len(audio_array)/self.sample_rate:.1f} 秒", fg='green')

    def import_audio(self):
        file_path = filedialog.askopenfilename(
            title="选择音频文件",
            filetypes=[("音频文件", "*.wav *.mp3 *.m4a *.ogg *.flac"), ("所有文件", "*.*")]
        )
        if file_path:
            self.audio_file = file_path
            self.audio_array = None
            self.btn_transcribe.config(state='normal')
            self.status_label.config(text=f"状态：✅ 已导入：{os.path.basename(file_path)}", fg='green')

    def import_word(self):
        file_path = filedialog.askopenfilename(
            title="选择手写记录文件",
            filetypes=[("Word文件", "*.docx"), ("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if not file_path:
            return

        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()

            if not text.strip():
                messagebox.showwarning("提示", "文件内容为空")
                return

            minutes = self.generate_minutes(text, source="手写记录整理")
            self.show_result(minutes)
            self.status_label.config(text=f"状态：✅ 已整理：{os.path.basename(file_path)}", fg='green')

        except Exception as e:
            messagebox.showerror("错误", f"读取文件失败：{str(e)}")

    def transcribe_and_generate(self):
        if not self.audio_file:
            messagebox.showwarning("提示", "请先录音或导入音频文件")
            return

        self.btn_transcribe.config(state='disabled')
        self.progress.start()
        self.status_label.config(text="状态：⏳ 正在加载语音识别模型...", fg='blue')

        def process():
            try:
                if self.model is None:
                    self.model = whisper.load_model("base")

                self.root.after(0, lambda: self.status_label.config(text="状态：⏳ 正在转写文字...", fg='blue'))

                if self.audio_array is not None:
                    audio_input = self.audio_array.flatten().astype('float32')
                    if self.sample_rate != 16000:
                        new_len = int(len(audio_input) / self.sample_rate * 16000)
                        audio_input = np.interp(
                            np.linspace(0, len(audio_input), new_len),
                            np.arange(len(audio_input)), audio_input
                        ).astype('float32')
                else:
                    data, sr = sf.read(self.audio_file, dtype='float32')
                    if data.ndim > 1:
                        data = data.mean(axis=1)
                    if sr != 16000:
                        new_len = int(len(data) / sr * 16000)
                        audio_input = np.interp(
                            np.linspace(0, len(data), new_len),
                            np.arange(len(data)), data
                        ).astype('float32')
                    else:
                        audio_input = data

                result = self.model.transcribe(audio_input, language='zh')
                transcript = result['text'].strip()
                minutes = self.generate_minutes(transcript, source="语音转写")
                self.root.after(0, lambda: self.show_result(minutes))

            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("错误", f"处理失败：{str(e)}"))
                self.root.after(0, lambda: self.btn_transcribe.config(state='normal'))
            finally:
                self.root.after(0, self.progress.stop)

        threading.Thread(target=process, daemon=True).start()

    def generate_minutes(self, content, source=""):
        now = datetime.datetime.now()
        date_str = now.strftime("%Y年%m月%d日 %H:%M")
        attendees = self.attendees_var.get() or "（未填写）"
        topic = self.topic_var.get() or "（未填写）"

        sentences = [s.strip() for s in content.replace('。', '。\n').replace('！', '！\n').replace('？', '？\n').replace('\n', '\n').split('\n') if s.strip()]

        minutes = f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        会 议 纪 要
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📅 会议时间：{date_str}
👥 参会人员：{attendees}
📌 会议主题：{topic}
📋 来源：{source}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【原始内容】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{content}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【要点整理】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

"""
        for i, sentence in enumerate(sentences[:15], 1):
            if sentence:
                minutes += f"{i}. {sentence}\n"

        minutes += f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【待办事项】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

（请手动补充待办事项）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
记录人：自动生成  |  生成时间：{date_str}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        return minutes

    def show_result(self, minutes):
        self.result_text.delete('1.0', tk.END)
        self.result_text.insert('1.0', minutes)
        self.btn_save_txt.config(state='normal')
        self.btn_save_word.config(state='normal')
        self.btn_transcribe.config(state='normal')
        self.status_label.config(text="状态：✅ 会议纪要生成完成！", fg='green')

    def save_minutes(self, fmt):
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        content = self.result_text.get('1.0', tk.END).strip()

        if fmt == 'txt':
            file_path = filedialog.asksaveasfilename(
                title="保存会议纪要", defaultextension=".txt",
                initialfile=f"会议纪要_{now}.txt",
                filetypes=[("文本文件", "*.txt")]
            )
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                messagebox.showinfo("成功", f"已保存到：\n{file_path}")

        elif fmt == 'docx':
            file_path = filedialog.asksaveasfilename(
                title="保存会议纪要", defaultextension=".docx",
                initialfile=f"会议纪要_{now}.docx",
                filetypes=[("Word文件", "*.docx")]
            )
            if file_path:
                doc = Document()
                doc.add_heading('会议纪要', 0)
                for line in content.split('\n'):
                    doc.add_paragraph(line)
                doc.save(file_path)
                messagebox.showinfo("成功", f"已保存到：\n{file_path}")


if __name__ == '__main__':
    root = tk.Tk()
    app = MeetingMinutesApp(root)
    root.mainloop()
